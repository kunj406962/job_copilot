import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from core.profile import Profile
from core.skills import load_skills

OUTPUT_DIR = "./output/resumes"


def _add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E75B6")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_section_heading(doc: Document, title: str) -> None:
    _add_horizontal_rule(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)


def _add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.lstrip("-• ").strip())
    run.font.size = Pt(10)


def _add_subsection_heading(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)


def _render_bulleted_section(doc: Document, section_text: str) -> None:
    if not section_text.strip():
        return
    for line in section_text.strip().split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith("-") or line.startswith("•"):
            _add_bullet(doc, line)
        else:
            _add_subsection_heading(doc, line)


def _render_skills_section(doc: Document) -> None:
    skills_data = load_skills()
    if not skills_data:
        return
    for category, skills in skills_data.items():
        if not skills:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        label_run = p.add_run(f"{category}: ")
        label_run.bold = True
        label_run.font.size = Pt(10)
        skills_run = p.add_run(", ".join(skills))
        skills_run.font.size = Pt(10)


def build_resume(
    profile: Profile,
    summary: str,
    projects: str,
    experience: str,
    filename: str,
) -> str:
    doc = Document()

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # --- Name ---
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run(profile.name)
    name_run.bold = True
    name_run.font.size = Pt(20)

    # --- Contact ---
    contact_parts = [x for x in [
        profile.email, profile.phone, profile.linkedin,
        profile.github, profile.location,
    ] if x]
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(6)
    contact_run = contact_p.add_run(" | ".join(contact_parts))
    contact_run.font.size = Pt(9)
    contact_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    # --- Summary ---
    _add_section_heading(doc, "Summary")
    summary_p = doc.add_paragraph()
    summary_p.paragraph_format.space_after = Pt(4)
    summary_run = summary_p.add_run(summary.strip())
    summary_run.font.size = Pt(10)

    # --- Experience ---
    if experience.strip():
        _add_section_heading(doc, "Experience")
        _render_bulleted_section(doc, experience)

    # --- Projects ---
    if projects.strip():
        _add_section_heading(doc, "Projects")
        _render_bulleted_section(doc, projects)

    # --- Skills (from skills.json) ---
    skills_data = load_skills()
    if skills_data:
        _add_section_heading(doc, "Skills")
        _render_skills_section(doc)

    # --- Education ---
    _add_section_heading(doc, "Education")
    for edu in profile.education:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(0)
        deg_run = p.add_run(f"{edu.degree} — {edu.institution}")
        deg_run.bold = True
        deg_run.font.size = Pt(10)

        detail_parts = [edu.graduation_year]
        if edu.gpa:
            detail_parts.append(f"GPA: {edu.gpa}")
        detail_p = doc.add_paragraph()
        detail_p.paragraph_format.space_before = Pt(0)
        detail_p.paragraph_format.space_after = Pt(4)
        detail_run = detail_p.add_run(" | ".join(detail_parts))
        detail_run.font.size = Pt(9)
        detail_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    output_path = os.path.join(OUTPUT_DIR, filename)
    doc.save(output_path)
    return output_path


def build_cover_letter(profile: Profile, cover_letter_text: str, filename: str) -> str:
    doc = Document()

    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run(profile.name)
    name_run.bold = True
    name_run.font.size = Pt(16)

    contact_parts = [x for x in [profile.email, profile.phone, profile.location] if x]
    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(20)
    contact_run = contact_p.add_run(" | ".join(contact_parts))
    contact_run.font.size = Pt(9)
    contact_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    for paragraph in cover_letter_text.strip().split("\n\n"):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        run = p.add_run(paragraph)
        run.font.size = Pt(10)

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    output_path = os.path.join(OUTPUT_DIR, filename)
    doc.save(output_path)
    return output_path
